import os
import pytesseract
import asyncio
import aiofiles

from pdf2image import convert_from_path
from docx import Document
from flask import Flask, request, redirect, url_for, render_template, send_from_directory
from pdf2docx import Converter




app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
CONVERT_FOLDER = 'convert'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CONVERT_FOLDER'] = CONVERT_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERT_FOLDER, exist_ok=True)

@app.route('/')
def index():
    options = ['Обычный pdf', 'Pdf OCR']
    return render_template('index.html', options=options)

@app.route('/convert', methods=['POST'])
def convert():
    selected_option = request.form['options']

    if 'pdfFile' not in request.files:
        return 'Нет файла для загрузки', 400

    file = request.files['pdfFile']
    
    if file.filename == '':
        return 'Нет выбранного файла', 400

    # Сохраняем файл перед конвертацией
    pdf_save_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(pdf_save_path)

    if selected_option == "Обычный pdf":
        return send_from_directory(app.config['CONVERT_FOLDER'], defaul_pdf(pdf_save_path), as_attachment=True)
    elif selected_option == "Pdf OCR":
        return send_from_directory(app.config['CONVERT_FOLDER'], ocr_pdf(pdf_save_path), as_attachment=True)

def ocr_pdf(pdf_file):
    images = convert_from_path(pdf_file)
    doc = Document()

    for image in images:
        text = pytesseract.image_to_string(image, lang='rus+eng')
        doc.add_paragraph(text)

    docx_file = f'{os.path.basename(pdf_file).split(".")[0]}.docx'
    docx_save_path = os.path.join(app.config['CONVERT_FOLDER'], docx_file)
    doc.save(docx_save_path)

    return docx_file

def defaul_pdf(pdf_file):
    docx_file = f'{os.path.basename(pdf_file).split(".")[0]}.docx'
    docx_save_path = os.path.join(app.config['CONVERT_FOLDER'], docx_file)

    cv = Converter(pdf_file)
    cv.convert(docx_save_path) 
    cv.close() 
    
    return docx_file


# async def ocr_pdf(pdf_file):
#     images = await asyncio.to_thread(convert_from_path, pdf_file)
#     doc = Document()

#     for image in images:
#         text = await asyncio.to_thread(pytesseract.image_to_string, image, lang='rus+eng')
#         doc.add_paragraph(text)

#     docx_file = f'{os.path.basename(pdf_file).split(".")[0]}.docx'
#     docx_save_path = os.path.join(app.config['CONVERT_FOLDER'], docx_file)
    
#     # Save the document asynchronously
#     async with aiofiles.open(docx_save_path, 'wb') as f:
#         await asyncio.to_thread(doc.save, f)

#     return docx_file

# async def defaul_pdf(pdf_file):
#     docx_file = f'{os.path.basename(pdf_file).split(".")[0]}.docx'
#     docx_save_path = os.path.join(app.config['CONVERT_FOLDER'], docx_file)

#     cv = Converter(pdf_file)
#     await asyncio.to_thread(cv.convert, docx_save_path)
#     await asyncio.to_thread(cv.close)

#     return docx_file


if __name__ == '__main__':
    app.run(host="172.17.41.75", port=5001)


